from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "generated_report"
DIAGRAM_DIR = OUT_DIR / "diagrams"
OUTPUT_DOCX = OUT_DIR / "Report4_Software Design Document - LMS Completed.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5A6673"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
GREEN = "E8F5E9"
YELLOW = "FFF7DF"
RED = "FDEBEC"
PURPLE = "F1EEFF"
WHITE = "FFFFFF"
BORDER = "A7B0BA"


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    fnt: ImageFont.FreeTypeFont,
    max_width: int,
) -> list[str]:
    words = text.replace("\n", " \n ").split()
    lines: list[str] = []
    current = ""
    for word in words:
        if word == "\n":
            if current:
                lines.append(current)
                current = ""
            continue
        candidate = word if not current else f"{current} {word}"
        if text_size(draw, candidate, fnt)[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.FreeTypeFont,
    fill: str = INK,
    line_gap: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, x2 - x1 - 28)
    total_h = sum(text_size(draw, line, fnt)[1] for line in lines) + line_gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        w, h = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=hex_to_rgb(fill))
        y += h + line_gap


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    subtitle: str = "",
    fill: str = WHITE,
    outline: str = BORDER,
    title_fill: str = INK,
    radius: int = 18,
) -> None:
    draw.rounded_rectangle(
        xy,
        radius=radius,
        fill=hex_to_rgb(fill),
        outline=hex_to_rgb(outline),
        width=3,
    )
    x1, y1, x2, y2 = xy
    title_font = font(34, True)
    body_font = font(29)
    if subtitle:
        title_h = text_size(draw, title, title_font)[1]
        draw_centered_text(draw, (x1 + 12, y1 + 16, x2 - 12, y1 + 58), title, title_font, title_fill)
        lines = wrap_text(draw, subtitle, body_font, x2 - x1 - 36)
        y = y1 + 72
        for line in lines[:4]:
            w, h = text_size(draw, line, body_font)
            draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=body_font, fill=hex_to_rgb(MUTED))
            y += h + 8
    else:
        draw_centered_text(draw, xy, title, title_font, title_fill)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = DARK_BLUE,
    width: int = 4,
) -> None:
    draw.line([start, end], fill=hex_to_rgb(fill), width=width)
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    size = 16
    left = (-uy, ux)
    p1 = (x2, y2)
    p2 = (int(x2 - ux * size + left[0] * size * 0.6), int(y2 - uy * size + left[1] * size * 0.6))
    p3 = (int(x2 - ux * size - left[0] * size * 0.6), int(y2 - uy * size - left[1] * size * 0.6))
    draw.polygon([p1, p2, p3], fill=hex_to_rgb(fill))


def canvas(path: Path, title: str, size: tuple[int, int] = (1800, 1100)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, hex_to_rgb("FFFFFF"))
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), title, font=font(38, True), fill=hex_to_rgb(INK))
    draw.line((60, 96, size[0] - 60, 96), fill=hex_to_rgb(BLUE_GRAY), width=4)
    return image, draw


def save(image: Image.Image, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=95)
    return path


def architecture_diagram() -> Path:
    path = DIAGRAM_DIR / "01_system_architecture.png"
    img, draw = canvas(path, "System Architecture - LMS")
    boxes = {
        "users": (70, 170, 380, 330, "Students / Teachers / Admins", "Browser users access SPA, REST API and Socket.IO"),
        "fe": (540, 150, 900, 330, "React 19 + Vite SPA", "Routes, pages, services, stores, hooks, UI components"),
        "be": (1080, 145, 1540, 345, "ExpressJS + TypeScript API", "REST routes, controllers, services, Mongoose models, Socket.IO server"),
        "db": (1050, 520, 1350, 670, "MongoDB + Mongoose", "31 domain collections and indexes"),
        "storage": (1430, 510, 1730, 670, "MinIO / S3 Storage", "Lesson files, submissions, avatars, chat attachments"),
        "email": (730, 520, 990, 670, "Resend Mail", "Verification, reset password, invites, absence warnings"),
        "google": (390, 520, 650, 670, "Google OAuth", "Optional sign-in identity provider"),
        "socket": (1060, 800, 1540, 980, "Socket.IO + WebRTC Signaling", "Chat rooms, typing, unread counts, video call offer/answer/ICE"),
    }
    fills = {
        "users": BLUE_GRAY,
        "fe": GREEN,
        "be": "E3F2FD",
        "db": YELLOW,
        "storage": "F3E5F5",
        "email": "FFF3E0",
        "google": "E8EAED",
        "socket": PURPLE,
    }
    for key, (x1, y1, x2, y2, title, subtitle) in boxes.items():
        draw_box(draw, (x1, y1, x2, y2), title, subtitle, fill=fills[key])
    arrow(draw, (380, 250), (540, 250))
    arrow(draw, (900, 245), (1080, 245))
    arrow(draw, (1300, 345), (1210, 520))
    arrow(draw, (1400, 345), (1540, 510))
    arrow(draw, (1080, 320), (990, 520))
    arrow(draw, (540, 330), (560, 520))
    arrow(draw, (1290, 345), (1280, 800))
    arrow(draw, (300, 330), (1060, 880))
    small = font(20, True)
    labels = [
        ((430, 220), "HTTPS UI"),
        ((950, 214), "Axios REST"),
        ((1230, 432), "CRUD / query"),
        ((1510, 430), "putObject / signed URL"),
        ((905, 430), "email API"),
        ((470, 430), "ID token"),
        ((1160, 700), "persistent socket"),
    ]
    for pos, label in labels:
        draw.text(pos, label, font=small, fill=hex_to_rgb(DARK_BLUE))
    return save(img, path)


def package_diagram_backend() -> Path:
    path = DIAGRAM_DIR / "02_backend_package_diagram.png"
    img, draw = canvas(path, "Backend Package Diagram - BE_LMS/src")
    layer_font = font(24, True)
    draw.text((80, 140), "Client-facing layer", font=layer_font, fill=hex_to_rgb(DARK_BLUE))
    draw.text((80, 365), "Business logic layer", font=layer_font, fill=hex_to_rgb(DARK_BLUE))
    draw.text((80, 590), "Data access layer", font=layer_font, fill=hex_to_rgb(DARK_BLUE))
    draw.text((80, 805), "Cross-cutting packages", font=layer_font, fill=hex_to_rgb(DARK_BLUE))
    boxes = {
        "routes": (170, 190, 470, 300, "routes", "API endpoint registration"),
        "socket": (560, 190, 860, 300, "socket", "Socket.IO event engines"),
        "controller": (280, 415, 640, 535, "controller", "HTTP orchestration"),
        "services": (780, 415, 1140, 535, "services", "Business rules and transactions"),
        "models": (610, 650, 970, 765, "models", "Mongoose schemas and indexes"),
        "middleware": (150, 860, 410, 965, "middleware", "auth, RBAC, errors"),
        "validators": (455, 860, 715, 965, "validators", "Zod request schemas"),
        "utils": (760, 860, 1020, 965, "utils", "JWT, mail, storage, asserts"),
        "config": (1065, 860, 1325, 965, "config", "DB, multer, MinIO, env"),
        "types": (1370, 860, 1630, 965, "types/constants", "interfaces, enums, codes"),
    }
    for key, (x1, y1, x2, y2, title, subtitle) in boxes.items():
        fill = GREEN if key in {"routes", "socket"} else "E3F2FD" if key in {"controller", "services"} else YELLOW if key == "models" else LIGHT_GRAY
        draw_box(draw, (x1, y1, x2, y2), title, subtitle, fill=fill)
    arrows = [
        ("routes", "controller"),
        ("routes", "middleware"),
        ("controller", "services"),
        ("controller", "validators"),
        ("services", "models"),
        ("services", "utils"),
        ("services", "config"),
        ("socket", "services"),
        ("socket", "utils"),
        ("models", "types"),
        ("middleware", "utils"),
        ("validators", "types"),
    ]
    centers = {k: ((v[0] + v[2]) // 2, (v[1] + v[3]) // 2) for k, v in boxes.items()}
    for a, b in arrows:
        start = centers[a]
        end = centers[b]
        arrow(draw, start, end, width=3)
    return save(img, path)


def package_diagram_frontend() -> Path:
    path = DIAGRAM_DIR / "03_frontend_package_diagram.png"
    img, draw = canvas(path, "Frontend Package Diagram - FE_LMS/src")
    boxes = {
        "routes": (90, 170, 390, 290, "routes", "React Router map"),
        "pages": (540, 170, 840, 290, "pages", "Feature screens"),
        "components": (990, 170, 1370, 290, "components", "Reusable UI blocks"),
        "services": (540, 420, 840, 540, "services", "Axios API clients"),
        "stores": (990, 420, 1370, 540, "stores/context", "auth, chat, socket, theme"),
        "hooks": (1410, 420, 1690, 540, "hooks", "stateful UI logic"),
        "types": (335, 675, 615, 795, "types", "DTOs and contracts"),
        "utils": (760, 675, 1040, 795, "utils", "http, date, slug, markdown"),
        "assets": (1185, 675, 1465, 795, "assets/public", "models, audio, textures"),
    }
    for key, (x1, y1, x2, y2, title, subtitle) in boxes.items():
        fill = GREEN if key in {"routes", "pages", "components"} else BLUE_GRAY if key in {"services", "stores", "hooks"} else LIGHT_GRAY
        draw_box(draw, (x1, y1, x2, y2), title, subtitle, fill=fill)
    centers = {k: ((v[0] + v[2]) // 2, (v[1] + v[3]) // 2) for k, v in boxes.items()}
    for a, b in [
        ("routes", "pages"),
        ("pages", "components"),
        ("pages", "services"),
        ("components", "stores"),
        ("components", "hooks"),
        ("services", "types"),
        ("services", "utils"),
        ("components", "assets"),
        ("stores", "services"),
    ]:
        arrow(draw, centers[a], centers[b], width=3)
    return save(img, path)


def erd_diagram(path_name: str, title: str, entities: list[tuple[str, str]], edges: list[tuple[int, int, str]]) -> Path:
    path = DIAGRAM_DIR / path_name
    img, draw = canvas(path, title)
    positions = [
        (100, 170, 410, 310),
        (535, 170, 845, 310),
        (970, 170, 1280, 310),
        (1405, 170, 1715, 310),
        (100, 455, 410, 595),
        (535, 455, 845, 595),
        (970, 455, 1280, 595),
        (1405, 455, 1715, 595),
        (315, 740, 625, 880),
        (755, 740, 1065, 880),
        (1195, 740, 1505, 880),
    ]
    centers = []
    for i, (name, desc) in enumerate(entities):
        x1, y1, x2, y2 = positions[i]
        draw_box(draw, (x1, y1, x2, y2), name, desc, fill=YELLOW if i % 2 else BLUE_GRAY)
        centers.append(((x1 + x2) // 2, (y1 + y2) // 2))
    rel_font = font(24, True)
    for a, b, label in edges:
        arrow(draw, centers[a], centers[b], width=3)
        mx = (centers[a][0] + centers[b][0]) // 2
        my = (centers[a][1] + centers[b][1]) // 2
        draw.rounded_rectangle((mx - 80, my - 20, mx + 80, my + 20), radius=8, fill=hex_to_rgb(WHITE), outline=hex_to_rgb(BLUE_GRAY))
        draw_centered_text(draw, (mx - 76, my - 18, mx + 76, my + 18), label, rel_font, DARK_BLUE, line_gap=2)
    return save(img, path)


def class_layer_diagram(path_name: str, title: str, classes: list[tuple[str, list[str], str]]) -> Path:
    path = DIAGRAM_DIR / path_name
    img, draw = canvas(path, title, size=(1800, 980))
    top = 170
    gap = 20
    box_w = 330
    x = 70
    centers = []
    for name, methods, fill in classes:
        h = 360
        xy = (x, top, x + box_w, top + h)
        draw.rounded_rectangle(xy, radius=16, fill=hex_to_rgb(fill), outline=hex_to_rgb(BORDER), width=3)
        draw.rectangle((x, top, x + box_w, top + 84), fill=hex_to_rgb(DARK_BLUE))
        draw_centered_text(draw, (x + 8, top + 8, x + box_w - 8, top + 78), name, font(24, True), WHITE)
        y = top + 101
        for method in methods[:5]:
            method_lines = wrap_text(draw, f"+ {method}", font(24), box_w - 38)
            for line in method_lines[:2]:
                draw.text((x + 20, y), line, font=font(24), fill=hex_to_rgb(INK))
                y += 28
            y += 6
        centers.append((x + box_w // 2, top + h // 2))
        x += box_w + gap
    for i in range(len(centers) - 1):
        arrow(draw, (centers[i][0] + box_w // 2 - 12, centers[i][1]), (centers[i + 1][0] - box_w // 2 + 12, centers[i + 1][1]), width=3)
    note = "Request flow: UI/Route -> Controller -> Service -> Model/External integration -> Response"
    draw_box(draw, (170, 760, 1630, 890), "Layered responsibility", note, fill=LIGHT_GRAY)
    return save(img, path)


def sequence_diagram(path_name: str, title: str, actors: list[str], steps: list[tuple[int, int, str]]) -> Path:
    path = DIAGRAM_DIR / path_name
    width = 1800
    height = max(1020, 340 + 145 * len(steps))
    img, draw = canvas(path, title, size=(width, height))
    left = 100
    right = width - 100
    y_top = 150
    lane_gap = (right - left) / (len(actors) - 1)
    xs = [int(left + i * lane_gap) for i in range(len(actors))]
    actor_font = font(26, True)
    label_font = font(29)
    for x, actor in zip(xs, actors):
        draw.rounded_rectangle((x - 110, y_top, x + 110, y_top + 58), radius=12, fill=hex_to_rgb(BLUE_GRAY), outline=hex_to_rgb(BORDER), width=2)
        draw_centered_text(draw, (x - 104, y_top + 6, x + 104, y_top + 52), actor, actor_font, INK)
        draw.line((x, y_top + 58, x, height - 70), fill=hex_to_rgb("D8DEE6"), width=3)
    y = y_top + 110
    for idx, (src, dst, label) in enumerate(steps, start=1):
        x1 = xs[src]
        x2 = xs[dst]
        y += 68
        if src == dst:
            draw.arc((x1 - 36, y - 22, x1 + 36, y + 42), 270, 90, fill=hex_to_rgb(DARK_BLUE), width=3)
            arrow(draw, (x1 + 36, y + 12), (x1 + 10, y + 35), width=3)
        else:
            arrow(draw, (x1, y), (x2, y), width=3)
        label_text = f"{idx}. {label}"
        max_w = abs(x2 - x1) - 30 if src != dst else 360
        max_w = max(max_w, 260)
        lx = min(x1, x2) + abs(x2 - x1) // 2 - max_w // 2 if src != dst else x1 + 50
        lines = wrap_text(draw, label_text, label_font, max_w)
        ly = y - 34
        for line in lines[:3]:
            draw.text((lx, ly), line, font=label_font, fill=hex_to_rgb(INK))
            ly += 34
        y += 24 * max(1, len(lines[:3]))
    return save(img, path)


def make_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "architecture": architecture_diagram(),
        "backend_package": package_diagram_backend(),
        "frontend_package": package_diagram_frontend(),
    }
    diagrams["erd_core"] = erd_diagram(
        "04_erd_core.png",
        "ERD - Core Academic Domain",
        [
            ("User", "student / teacher / admin"),
            ("Specialist", "teaching specialization"),
            ("Major", "academic major"),
            ("Subject", "catalog + prerequisites"),
            ("Semester", "academic term"),
            ("Course", "class section"),
            ("Enrollment", "student-course status"),
            ("CourseInvite", "secure join token"),
        ],
        [(1, 2, "N:1"), (3, 1, "N:N"), (5, 3, "N:1"), (5, 4, "N:1"), (6, 0, "N:1"), (6, 5, "N:1"), (7, 5, "N:1")],
    )
    diagrams["erd_learning"] = erd_diagram(
        "05_erd_learning_assessment.png",
        "ERD - Learning Content and Assessment",
        [
            ("Course", "course shell"),
            ("Lesson", "ordered module"),
            ("LessonMaterial", "file metadata + S3 key"),
            ("LessonProgress", "time and completion"),
            ("Quiz", "test settings"),
            ("QuizQuestion", "question bank"),
            ("QuizAttempt", "answers + score"),
            ("Assignment", "essay/file task"),
            ("Submission", "student upload + grade"),
            ("User", "student / teacher"),
        ],
        [(1, 0, "N:1"), (2, 1, "N:1"), (3, 1, "N:1"), (3, 9, "N:1"), (4, 0, "N:1"), (6, 4, "N:1"), (6, 9, "N:1"), (7, 0, "N:1"), (8, 7, "N:1"), (8, 9, "N:1"), (5, 0, "by subject")],
    )
    diagrams["erd_schedule"] = erd_diagram(
        "06_erd_schedule_attendance.png",
        "ERD - Schedule and Attendance",
        [
            ("Course", "class section"),
            ("ClassSchedule", "weekly plan"),
            ("TimeSlot", "start/end slot"),
            ("ScheduleException", "make-up/cancel/change"),
            ("Attendance", "per student per date"),
            ("User", "student / teacher / admin"),
            ("Notification", "system notice"),
        ],
        [(1, 0, "N:1"), (1, 2, "N:1"), (1, 5, "teacher"), (3, 1, "N:1"), (3, 2, "new slot"), (4, 0, "N:1"), (4, 5, "student"), (6, 5, "recipient")],
    )
    diagrams["erd_communication"] = erd_diagram(
        "07_erd_communication.png",
        "ERD - Communication and Support",
        [
            ("Course", "course context"),
            ("Forum", "discussion board"),
            ("ForumPost", "topic"),
            ("ForumReply", "nested reply"),
            ("ChatRoom", "direct/group room"),
            ("Message", "text/file/system msg"),
            ("Announcement", "course/system bulletin"),
            ("Feedback", "rating/review"),
            ("Blog", "public article"),
            ("User", "author/participant"),
        ],
        [(1, 0, "N:1"), (2, 1, "N:1"), (2, 9, "author"), (3, 2, "N:1"), (3, 9, "author"), (4, 0, "optional"), (5, 4, "N:1"), (5, 9, "sender"), (6, 0, "optional"), (6, 9, "author"), (7, 9, "author"), (8, 9, "author")],
    )

    module_classes = {
        "auth": [
            ("AuthPage/AuthRoute", ["login()", "register()", "resetPassword()"], GREEN),
            ("auth.route.ts", ["POST /login", "POST /register", "POST /refresh"], BLUE_GRAY),
            ("AuthController", ["loginHandler()", "registerHandler()", "refreshHandler()"], "E3F2FD"),
            ("AuthService", ["loginUser()", "registerUser()", "refreshAccessToken()"], "E3F2FD"),
            ("Auth Models", ["User.comparePassword()", "Session record", "VerificationCode OTP"], YELLOW),
        ],
        "enrollment": [
            ("CourseDetail UI", ["selfEnroll()", "joinInvite()"], GREEN),
            ("enrollment.route.ts", ["POST /enroll", "PUT /:id", "POST /:id/kick"], BLUE_GRAY),
            ("EnrollmentController", ["create", "update", "statistics"], "E3F2FD"),
            ("EnrollmentService", ["check prerequisites", "cooldown", "capacity"], "E3F2FD"),
            ("Academic Models", ["Course ongoing only", "unique student-course", "status lifecycle"], YELLOW),
        ],
        "lesson": [
            ("Lesson UI", ["view material", "upload file", "track time"], GREEN),
            ("lesson-materials route", ["POST /upload", "GET /:id/download"], BLUE_GRAY),
            ("LessonMaterialController", ["upload", "get", "delete file"], "E3F2FD"),
            ("LessonMaterialService", ["access control", "presigned URL", "MinIO key"], "E3F2FD"),
            ("Lesson/Material/Progress", ["timeSpentSeconds", "lastAccessedAt", "isCompleted"], YELLOW),
        ],
        "quiz": [
            ("TakeQuizPage", ["start", "auto-save", "submit"], GREEN),
            ("quiz-attempt.route.ts", ["POST /enroll", "PUT /auto-save", "PUT /submit"], BLUE_GRAY),
            ("QuizAttemptController", ["enroll", "autoSave", "ban", "re-grade"], "E3F2FD"),
            ("QuizAttemptService", ["validate student", "grade answers", "lock status"], "E3F2FD"),
            ("Quiz/QuizAttempt", ["snapshotQuestions", "answers", "score/status"], YELLOW),
        ],
        "assignment": [
            ("Assignment UI", ["create task", "submit file", "grade"], GREEN),
            ("submission.route.ts", ["POST /", "PUT /grade", "GET /stats"], BLUE_GRAY),
            ("SubmissionController", ["submit", "grade", "report"], "E3F2FD"),
            ("SubmissionService", ["late check", "MinIO file", "grade history"], "E3F2FD"),
            ("Assignment/Submission", ["dueDate", "maxScore", "feedback"], YELLOW),
        ],
        "attendance": [
            ("AttendancePage", ["mark", "export", "send warning"], GREEN),
            ("attendance.route.ts", ["POST /", "PATCH /:id", "send notifications"], BLUE_GRAY),
            ("AttendanceController", ["create", "update", "sendAbsence"], "E3F2FD"),
            ("AttendanceService", ["stats", "threshold 20%", "email batch"], "E3F2FD"),
            ("Attendance/Schedule", ["unique course-student-date", "status"], YELLOW),
        ],
        "chat": [
            ("Chat UI", ["message", "file", "video call"], GREEN),
            ("Socket.IO Gateway", ["send-message", "typing", "video offer"], BLUE_GRAY),
            ("socketConversation", ["create message", "unread count", "notify room"], "E3F2FD"),
            ("socketVideoCall", ["activeCalls", "offer/answer", "ICE candidate"], "E3F2FD"),
            ("ChatRoom/Message", ["participants", "lastMessage", "file metadata"], YELLOW),
        ],
        "feedback": [
            ("Feedback/Announcement UI", ["create", "list", "moderate"], GREEN),
            ("API routes", ["POST /", "GET /", "DELETE /:id"], BLUE_GRAY),
            ("Controllers", ["validate request", "authorize owner"], "E3F2FD"),
            ("Services", ["rating aggregation", "target filters", "notifications"], "E3F2FD"),
            ("Comm Models", ["Feedback target", "Announcement scope", "Notification isRead"], YELLOW),
        ],
    }
    module_titles = {
        "auth": "Class Diagram - Authentication and Session",
        "enrollment": "Class Diagram - Course Enrollment",
        "lesson": "Class Diagram - Lesson Material and Progress",
        "quiz": "Class Diagram - Quiz Attempt Lifecycle",
        "assignment": "Class Diagram - Assignment Submission",
        "attendance": "Class Diagram - Attendance Alerts",
        "chat": "Class Diagram - Chat and Video Call",
        "feedback": "Class Diagram - Feedback and Announcements",
    }
    for key, classes in module_classes.items():
        diagrams[f"class_{key}"] = class_layer_diagram(f"10_class_{key}.png", module_titles[key], classes)

    diagrams["seq_auth"] = sequence_diagram(
        "20_seq_auth.png",
        "Sequence Diagram - Login / Register / Refresh",
        ["User", "React UI", "Auth API", "AuthService", "MongoDB", "Resend"],
        [
            (0, 1, "Enter credentials or registration data"),
            (1, 2, "POST /auth/login or /auth/register"),
            (2, 3, "Validate DTO and call service"),
            (3, 4, "Find/create User; create Session/VerificationCode"),
            (3, 5, "Send verification or reset email when needed"),
            (3, 2, "Return accessToken + refreshToken"),
            (2, 1, "Set HTTP-only cookies"),
            (1, 0, "Navigate to role dashboard"),
        ],
    )
    diagrams["seq_enrollment"] = sequence_diagram(
        "21_seq_enrollment.png",
        "Sequence Diagram - Student Self Enrollment",
        ["Student", "Course UI", "Enrollment API", "EnrollmentService", "MongoDB", "Notification"],
        [
            (0, 1, "Open course detail and click Enroll"),
            (1, 2, "POST /enrollments/enroll"),
            (2, 3, "Authorize role=student and parse payload"),
            (3, 4, "Load Course, Subject prerequisites, existing Enrollment"),
            (3, 3, "Check ongoing status, password, cooldown, capacity"),
            (3, 4, "Create or reactivate Enrollment"),
            (3, 5, "Notify teacher if approval is required"),
            (2, 1, "Return enrollment status"),
        ],
    )
    diagrams["seq_lesson"] = sequence_diagram(
        "22_seq_lesson.png",
        "Sequence Diagram - Lesson Material Access and Progress",
        ["Student/Teacher", "React UI", "Lesson API", "LessonService", "MongoDB", "MinIO"],
        [
            (0, 1, "Open lesson or upload material"),
            (1, 2, "GET /lesson-materials or POST /upload"),
            (2, 3, "Check role and course access"),
            (3, 4, "Read Lesson, Course, Enrollment, Material metadata"),
            (3, 5, "Upload file or generate presigned URL"),
            (3, 4, "Persist LessonMaterial or LessonProgress"),
            (2, 1, "Return material URL / progress data"),
        ],
    )
    diagrams["seq_quiz"] = sequence_diagram(
        "23_seq_quiz.png",
        "Sequence Diagram - Quiz Attempt Auto-save and Submit",
        ["Student", "TakeQuizPage", "QuizAttempt API", "QuizAttemptService", "MongoDB", "Teacher"],
        [
            (0, 1, "Start quiz"),
            (1, 2, "POST /quiz-attempts/enroll"),
            (2, 3, "Validate enrollment and status"),
            (3, 4, "Create or return in_progress attempt"),
            (1, 2, "PUT /:id/auto-save as answers change"),
            (3, 4, "Store answers without final score"),
            (1, 2, "PUT /:id/submit"),
            (3, 4, "Grade answers, set submitted, score"),
            (5, 2, "Optional ban or re-grade"),
        ],
    )
    diagrams["seq_assignment"] = sequence_diagram(
        "24_seq_assignment.png",
        "Sequence Diagram - Assignment Submission and Grading",
        ["Student", "Assignment UI", "Submission API", "SubmissionService", "MongoDB", "MinIO", "Teacher"],
        [
            (0, 1, "Choose assignment and file"),
            (1, 2, "POST /submissions"),
            (2, 3, "Validate ownership and deadline"),
            (3, 5, "Upload file buffer"),
            (3, 4, "Create/update Submission"),
            (6, 2, "PUT /submissions/:id/grade"),
            (3, 4, "Save grade, feedback, gradeHistory"),
            (2, 1, "Return grade/status"),
        ],
    )
    diagrams["seq_attendance"] = sequence_diagram(
        "25_seq_attendance.png",
        "Sequence Diagram - Attendance and Absence Warning",
        ["Teacher", "Attendance UI", "Attendance API", "AttendanceService", "MongoDB", "Resend"],
        [
            (0, 1, "Open course roster"),
            (1, 2, "POST /attendances or PATCH /:id"),
            (2, 3, "Authorize teacher/admin"),
            (3, 4, "Upsert course-student-date status"),
            (0, 1, "Trigger absence warning"),
            (1, 2, "POST /courses/:id/send-absence-notifications"),
            (3, 4, "Calculate absent percentage threshold 20%"),
            (3, 5, "Send warning emails"),
        ],
    )
    diagrams["seq_chat"] = sequence_diagram(
        "26_seq_chat.png",
        "Sequence Diagram - Chat Message and Video Call Signaling",
        ["User A", "Socket.IO", "Chat Service", "MongoDB", "User B", "WebRTC"],
        [
            (0, 1, "Connect with accessToken cookie"),
            (1, 2, "Authenticate socket and join rooms"),
            (0, 1, "chatroom:send-message"),
            (2, 3, "Create Message and update lastMessage/unreadCounts"),
            (1, 4, "Emit chatroom:new-message"),
            (0, 1, "videocall:start / offer"),
            (1, 4, "Emit incoming call / offer"),
            (4, 5, "Peer-to-peer media after offer/answer/ICE"),
        ],
    )
    diagrams["seq_feedback"] = sequence_diagram(
        "27_seq_feedback.png",
        "Sequence Diagram - Feedback and Announcement",
        ["User", "React UI", "API Route", "Service", "MongoDB", "Notification"],
        [
            (0, 1, "Create feedback or announcement"),
            (1, 2, "POST /feedbacks or /announcements"),
            (2, 3, "Validate role and payload"),
            (3, 4, "Persist target/rating/content"),
            (3, 5, "Create notification when applicable"),
            (2, 1, "Return saved record"),
            (0, 1, "View list/filter/detail/delete"),
            (1, 2, "GET/DELETE endpoint with ownership rule"),
        ],
    )
    return diagrams


def set_font(run, name: str = "Calibri", size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))


def set_style_font(style, name: str, size: int, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor(*hex_to_rgb(color))
    if bold is not None:
        style.font.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], "Calibri", 11, "000000")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    set_style_font(styles["Heading 1"], "Calibri", 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], "Calibri", 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], "Calibri", 12, DARK_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    for style_name in ("List Bullet", "List Number"):
        set_style_font(styles[style_name], "Calibri", 11, "000000")
        styles[style_name].paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("LMS Software Design Document | Report 4")
    set_font(run, size=9, color=MUTED)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
                    tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                    tcw = tc_pr.find(qn("w:tcW"))
                    if tcw is None:
                        tcw = OxmlElement("w:tcW")
                        tc_pr.append(tcw)
                    tcw.set(qn("w:w"), str(int(width * 1440)))
                    tcw.set(qn("w:type"), "dxa")


def format_table(table, widths: list[float] | None = None, header_fill: str = LIGHT_GRAY, font_size: int = 9) -> None:
    set_table_width(table, widths)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margin(cell)
            if ri == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run, size=font_size, bold=(ri == 0), color=INK)


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        rest = text[len(bold_prefix) :]
        if rest:
            r = p.add_run(rest)
            set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED, bold=False)
    r.italic = True


def add_picture(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: int = 9) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    format_table(table, widths, font_size=font_size)
    doc.add_paragraph()


BACKEND_PACKAGES = [
    ["01", "routes", "Đăng ký endpoint REST như /auth, /courses, /enrollments, /quiz-attempts; gắn middleware authenticate/authorize và chuyển request đến controller."],
    ["02", "controller", "Nhận request/response Express, parse dữ liệu qua validator, lấy user context từ middleware và gọi service tương ứng."],
    ["03", "services", "Chứa nghiệp vụ lõi: kiểm tra prerequisite, capacity, cooldown, grading, attendance threshold, soft delete và transaction."],
    ["04", "models", "Định nghĩa Mongoose schema, collection, ref, index và hook cho 31 entity nghiệp vụ."],
    ["05", "middleware", "Xác thực JWT, phân quyền RBAC, response helper, error handler tập trung."],
    ["06", "validators", "Schema kiểm tra request theo từng module để chặn dữ liệu sai trước khi vào service."],
    ["07", "socket", "Khởi tạo Socket.IO, xác thực socket qua cookie, xử lý chat room, typing, seen state và WebRTC signaling."],
    ["08", "config", "Cấu hình MongoDB, Multer memory storage, MinIO/S3 client, Resend và biến môi trường."],
    ["09", "utils", "Hàm dùng chung: JWT, bcrypt, sendMail, uploadFile, presigned URL, prefix file, appAssert/AppError."],
    ["10", "types", "Interface, enum và DTO TypeScript: Role, CourseStatus, EnrollmentStatus, AttemptStatus, AttendanceStatus..."],
    ["11", "constants", "HTTP status, mã lỗi ứng dụng, regex, giới hạn field, env constants."],
]


FRONTEND_PACKAGES = [
    ["01", "routes", "Khai báo React Router, ProtectedRoute và phân nhánh màn hình theo student/teacher/admin."],
    ["02", "pages", "Các màn hình nghiệp vụ: dashboard, courses, quizzes, assignments, attendance, forum, chat, feedback, curriculum."],
    ["03", "components", "UI tái sử dụng: layout, course tabs, quiz cards, calendar, chat window, video call, lesson material viewer."],
    ["04", "services", "Axios API client theo module: auth, course, enrollment, quiz, submission, lesson material, chat room..."],
    ["05", "stores/context", "State chung bằng Zustand/Context: auth, chat rooms, floating chat, video call, socket, theme."],
    ["06", "hooks", "Hook điều phối message, unread chat, video call, debounce, auth và theme."],
    ["07", "types", "TypeScript contract dùng ở FE cho Course, Quiz, Assignment, Schedule, Feedback, Auth..."],
    ["08", "utils", "Tiện ích HTTP, date, slug, markdown sanitize/render."],
    ["09", "assets/public", "Texture, model, audio và tài nguyên tĩnh dùng cho trải nghiệm học tập."],
]


TABLE_DESCRIPTIONS = [
    ["01", "User", "_id", "specialistIds -> Specialist", "Tài khoản Student/Teacher/Admin; lưu email, mật khẩu bcrypt, role, trạng thái, hồ sơ, Google ID."],
    ["02", "Session", "_id", "userId -> User", "Phiên đăng nhập dùng cho refresh token, user agent và thời hạn session."],
    ["03", "VerificationCode", "_id", "userId -> User", "Mã tạm thời cho verify email và reset password, có email/type/expiresAt."],
    ["04", "Specialist", "_id", "majorId -> Major", "Chuyên ngành dùng để ràng buộc giáo viên với Subject/Course phù hợp."],
    ["05", "Major", "_id", "-", "Ngành học hoặc khoa, có name/slug/description và timestamp."],
    ["06", "Subject", "_id", "specialistIds -> Specialist; prerequisites -> Subject", "Môn học chuẩn trong catalog; hỗ trợ prerequisite map, code, credits, active flag."],
    ["07", "Semester", "_id", "-", "Học kỳ/năm học với startDate/endDate; ràng buộc unique theo name và year/type."],
    ["08", "Course", "_id", "semesterId -> Semester; subjectId -> Subject; teacherIds/createdBy/deletedBy -> User", "Lớp học cụ thể của một subject; có status draft/ongoing/completed, capacity, publish flag, weight, soft delete."],
    ["09", "CourseInvite", "_id", "courseId -> Course; createdBy -> User", "Token join course có hash, expiry, maxUses, usedCount, active/deleted flags."],
    ["10", "Enrollment", "_id", "studentId/respondedBy -> User; courseId -> Course", "Bản ghi student tham gia course; trạng thái pending/approved/rejected/cancelled/dropped/completed và progress tổng hợp."],
    ["11", "Lesson", "_id", "courseId -> Course; createdBy -> User", "Bài học trong course, có thứ tự, nội dung, duration, publish status."],
    ["12", "LessonMaterial", "_id", "lessonId -> Lesson; uploadedBy -> User", "Metadata tài liệu: title, note, originalName, mimeType, key MinIO, size."],
    ["13", "LessonProgress", "_id", "lessonId -> Lesson; courseId -> Course; studentId -> User", "Theo dõi completedAt, timeSpentSeconds, lastAccessedAt cho từng student/lesson."],
    ["14", "Quiz", "_id", "courseId -> Course; createdBy -> User", "Cấu hình quiz: title, time window, shuffleQuestions, snapshotQuestions, published flag."],
    ["15", "QuizQuestion", "_id", "subjectId -> Subject", "Ngân hàng câu hỏi; lưu text, images, options, correctOptions, points, explanation."],
    ["16", "QuizAttempt", "_id", "quizId -> Quiz; studentId -> User", "Lượt làm bài; lưu startedAt, submittedAt, duration, answers, score, status."],
    ["17", "Assignment", "_id", "courseId -> Course; createdBy -> User", "Bài tập tự luận/file upload; có maxScore, dueDate, allowLate và file đính kèm của đề."],
    ["18", "Submission", "_id", "assignmentId -> Assignment; studentId/gradedBy -> User", "Bài nộp của student; lưu file key, trạng thái, điểm, feedback, gradeHistory, isLate."],
    ["19", "ClassSchedule", "_id", "courseId -> Course; teacherId/requestedBy/approvedBy -> User; timeSlotId -> TimeSlot", "Lịch học tuần; model name là ClassSchedule, collection classschedules; có status và recurrencePattern."],
    ["20", "TimeSlot", "_id", "createdBy -> User", "Ca học chuẩn: slotNumber, startTime, endTime, duration, applicableDays."],
    ["21", "ScheduleException", "_id", "scheduleId -> ClassSchedule; courseId -> Course; newTimeSlotId -> TimeSlot; replacementTeacherId/requestedBy/approvedBy -> User", "Ngoại lệ lịch học: hủy, học bù, đổi phòng, đổi giáo viên, trạng thái duyệt."],
    ["22", "Attendance", "_id", "courseId -> Course; studentId/markedBy -> User", "Điểm danh theo course-student-date, status notyet/present/absent; unique theo courseId/studentId/date."],
    ["23", "Forum", "_id", "courseId -> Course; createdBy -> User", "Không gian thảo luận theo course; có title/description và text index."],
    ["24", "ForumPost", "_id", "forumId -> Forum; authorId -> User", "Chủ đề thảo luận trong forum, hỗ trợ pinned và index theo forum/author."],
    ["25", "ForumReply", "_id", "postId -> ForumPost; authorId -> User; parentReplyId -> ForumReply", "Reply phân cấp cho forum post."],
    ["26", "ChatRoom", "_id", "courseId -> Course; participants.userId/createdBy/seenBy -> User; lastMessage.id -> Message", "Phòng chat trực tiếp/nhóm; lưu participants, lastMessage, unreadCounts."],
    ["27", "Message", "_id", "chatRoomId -> ChatRoom; senderId -> User", "Tin nhắn text/link/file/system trong chat room; có metadata file khi gửi đính kèm."],
    ["28", "Notification", "_id", "sender/recipientUser -> User; recipientCourse -> Course", "Thông báo trong hệ thống; có recipientType, isRead, soft delete."],
    ["29", "Announcement", "_id", "courseId -> Course; authorId -> User", "Thông báo course/system với title/content/publishedAt."],
    ["30", "Blog", "_id", "-", "Bài viết public; có title, slug unique, content, thumbnailUrl, authorName, avatar."],
    ["31", "Feedback", "_id", "userId -> User; targetId -> dynamic target", "Đánh giá/rating và phản hồi có thể kèm file minh chứng; targetId dùng cho course/teacher/system tùy type."],
]


FEATURES = [
    {
        "key": "auth",
        "heading": "3.1 Authentication & Session Management",
        "purpose": "Đảm bảo người dùng đăng ký, xác thực email, đăng nhập, refresh token, logout và reset password bằng luồng an toàn dựa trên JWT cookie + SessionModel.",
        "classes": [
            ["AuthController", "registerHandler, loginHandler, refreshHandler, verifyEmailHandler, resetPasswordHandler"],
            ["AuthService", "registerUser, loginUser, refreshUserAccessToken, verifyEmail, resetPassword, googleLogin"],
            ["UserModel", "So sánh mật khẩu bằng comparePassword(), lưu role/status/isVerified/specialistIds."],
            ["SessionModel", "Lưu sessionId trong refresh token, expiresAt và userAgent."],
            ["VerificationCodeModel", "Lưu OTP/link verify email hoặc forgot password."],
        ],
        "flow": [
            "Register: FE gửi /auth/register, service kiểm tra email trùng, tạo User, tạo VerificationCode và gửi email qua Resend.",
            "Login: service tìm user theo email, so sánh bcrypt, tạo Session, ký accessToken/refreshToken và controller set HTTP-only cookies.",
            "Refresh: controller đọc refreshToken cookie, service xác thực session còn hạn, gia hạn session nếu gần hết hạn và trả accessToken mới.",
            "Reset password: service xác thực VerificationCode loại FORGOT_PASSWORD, hash mật khẩu mới, xóa session cũ để buộc đăng nhập lại.",
        ],
    },
    {
        "key": "enrollment",
        "heading": "3.2 Course, Curriculum & Enrollment",
        "purpose": "Thiết kế luồng tạo course, publish course, self-enroll và quản lý trạng thái học tập dựa trên CourseStatus, EnrollmentStatus và prerequisite map của Subject.",
        "classes": [
            ["CourseController/Service", "list, create, update, soft delete, restore, complete course, statistics"],
            ["EnrollmentController/Service", "self enroll, admin enroll, update status, kick student, enrollment statistics"],
            ["SubjectModel", "Quản lý code/slug/specialistIds/prerequisites."],
            ["CourseModel", "semesterId, subjectId, teacherIds, capacity, weight, status, isPublished, soft delete."],
            ["EnrollmentModel", "studentId + courseId unique, status lifecycle, progress and finalGrade."],
        ],
        "flow": [
            "Teacher/Admin tạo Course ở draft; service bắt buộc có teacherIds, capacity hợp lệ, teacher active và specialist của teacher phải match specialist của Subject.",
            "Admin publish course: isPublished=true sẽ chuyển status từ draft sang ongoing và thông báo cho teacher.",
            "Student self-enroll: chỉ nhận course ongoing, kiểm tra prerequisite đã completed, password nếu có, cooldown 1 phút khi re-enroll và capacity còn chỗ.",
            "Teacher/Admin cập nhật enrollment: approval/rejection/completed/dropped tạo timestamp tương ứng và gửi Notification cho student.",
        ],
    },
    {
        "key": "lesson",
        "heading": "3.3 Lessons, Materials & Learning Progress",
        "purpose": "Quản lý cấu trúc bài học, tài liệu học tập trên MinIO/S3 và tiến độ học tập cá nhân của student.",
        "classes": [
            ["LessonController/Service", "create/update/delete lesson, list by course, enforce order."],
            ["LessonMaterialService", "role-based access, uploadFile, removeFile, presigned URL."],
            ["LessonProgressService", "add time, mark complete, get lesson/course progress."],
            ["MinIO Client", "putObject, removeObject, presignedGetObject."],
            ["Lesson/LessonMaterial/LessonProgress", "metadata and tracking models."],
        ],
        "flow": [
            "Teacher upload tài liệu qua /lesson-materials/upload; Multer memoryStorage kiểm tra MIME và limit 20MB.",
            "Service xác nhận user là instructor của course, upload file lên MinIO với prefix theo course/lesson, lưu key/originalName/mimeType/size.",
            "Student mở bài học; service xác nhận enrollment, đọc material metadata và tạo presigned URL để tải/stream file.",
            "FE gửi heartbeat PATCH /lesson-progress/lessons/:lessonId/time; service cộng timeSpentSeconds và cập nhật lastAccessedAt.",
        ],
    },
    {
        "key": "quiz",
        "heading": "3.4 Quiz, Question Bank & Attempt Lifecycle",
        "purpose": "Thiết kế ngân hàng câu hỏi, quiz được publish theo course, lượt làm bài có auto-save, submit, ban và re-grade.",
        "classes": [
            ["QuizQuestionService", "import/export XML, random question selection, CRUD question bank."],
            ["QuizService", "create/update/delete quiz, snapshotQuestions, statistics."],
            ["QuizAttemptService", "enroll, save, autoSave, submit, ban, re-grade."],
            ["QuizAttemptModel", "answers, score, status in_progress/submitted/abandoned."],
            ["QuizModel", "time window, shuffleQuestions, hashPassword, isPublished."],
        ],
        "flow": [
            "Teacher tạo quiz từ bank hoặc snapshotQuestions; route /quizzes yêu cầu role teacher/admin.",
            "Student start quiz qua /quiz-attempts/enroll; service xác nhận enrollment approved và không bị abandoned/submitted.",
            "Trong lúc làm bài, FE gọi /quiz-attempts/:id/auto-save; service ghi answers nháp để chống mất dữ liệu.",
            "Submit: service grade answers, cập nhật score/status=submitted/submittedAt; teacher/admin có thể ban hoặc re-grade khi có sự cố.",
        ],
    },
    {
        "key": "assignment",
        "heading": "3.5 Assignments, Submissions & Grading",
        "purpose": "Hỗ trợ teacher tạo bài tập, student nộp file, teacher chấm điểm và xem thống kê/report.",
        "classes": [
            ["AssignmentService", "CRUD assignment, course scope, deadline/maxScore."],
            ["SubmissionService", "submit/resubmit, status, late check, grade and reports."],
            ["UploadFile utility", "MinIO upload and metadata extraction."],
            ["AssignmentModel", "courseId, title, maxScore, dueDate, allowLate, file metadata."],
            ["SubmissionModel", "assignmentId, studentId, file key, grade, feedback, gradeHistory."],
        ],
        "flow": [
            "Teacher tạo Assignment cho course, có dueDate, maxScore và optional file đề bài.",
            "Student nộp bài qua /submissions; service kiểm tra quyền học course, deadline/allowLate, upload file lên MinIO và lưu Submission.",
            "Teacher chấm qua /submissions/:assignmentId/grade hoặc /by-submission/:id/grade; service lưu grade, feedback và gradeHistory.",
            "Teacher/Admin xem stats/report theo assignment hoặc course để tổng hợp điểm.",
        ],
    },
    {
        "key": "attendance",
        "heading": "3.6 Schedule & Attendance",
        "purpose": "Quản lý lịch học, ca học, ngoại lệ lịch và điểm danh; tự động cảnh báo vắng khi vượt ngưỡng.",
        "classes": [
            ["ScheduleService", "time slots, teacher availability, pending approve, exceptions."],
            ["AttendanceService", "create/update/delete/export, course/student stats, absence warning."],
            ["TimeSlotModel", "slotNumber, start/end, applicableDays."],
            ["ClassScheduleModel", "courseId, teacherId, dayOfWeek, timeSlotId, status."],
            ["AttendanceModel", "courseId, studentId, date, status, markedBy."],
        ],
        "flow": [
            "Teacher/Admin tạo schedule request; service kiểm tra availability và unique active teacher-day-slot.",
            "Admin approve schedule hoặc schedule exception để chuyển trạng thái hợp lệ.",
            "Teacher mark attendance theo course/date; service lưu unique courseId+studentId+date.",
            "Khi gửi cảnh báo, service tính số buổi absent; threshold mặc định 20% và gửi email qua Resend cho học viên liên quan.",
        ],
    },
    {
        "key": "chat",
        "heading": "3.7 Forums, Chat Rooms & WebRTC Video Calls",
        "purpose": "Thiết kế tương tác realtime: forum theo course, chat room, tin nhắn file, typing/unread và WebRTC signaling.",
        "classes": [
            ["ForumController/Service", "forum, post, reply CRUD with course context."],
            ["SocketAuthMiddleware", "Đọc accessToken từ cookie, verify JWT, attach socket.user."],
            ["socketConversation", "send message/file, mark as read, invite/leave/join room."],
            ["socketVideoCall", "activeCalls map, start/join/leave/offer/answer/ICE/end."],
            ["ChatRoom/Message", "participants, unreadCounts, lastMessage, file metadata."],
        ],
        "flow": [
            "Socket connection được xác thực bằng accessToken cookie; user join room riêng của mình và các chatRoom đang tham gia.",
            "Gửi message: socketConversation tạo Message, cập nhật ChatRoom.lastMessage/unreadCounts và emit chatroom:new-message.",
            "Gửi file chat: file buffer được upload lên MinIO, Message lưu file metadata rồi broadcast về room.",
            "Video call: socketVideoCall tạo callId, quản lý activeCalls/userCalls và relay offer/answer/ICE để client thiết lập WebRTC peer-to-peer.",
        ],
    },
    {
        "key": "feedback",
        "heading": "3.8 Feedback, Announcements, Notifications & Blog",
        "purpose": "Bao phủ các luồng phản hồi, thông báo khóa học/hệ thống, notification đọc/chưa đọc và bài viết public.",
        "classes": [
            ["FeedbackController/Service", "create, list by target, my-feedbacks, delete moderation."],
            ["AnnouncementController/Service", "course/system announcement CRUD."],
            ["NotificationService", "create notification, read/read-all, soft/hard delete."],
            ["BlogController/Service", "public CRUD by slug."],
            ["Feedback/Announcement/Notification/Blog", "supporting communication models."],
        ],
        "flow": [
            "Student gửi feedback có title, description, rating và optional attachment; service lưu targetId/type và file metadata nếu có.",
            "Teacher/Admin tạo announcement; route yêu cầu authenticate và service lưu course/system scope.",
            "Notification được tạo bởi các service khác khi enrollment/course/chat thay đổi; user đọc qua /notifications và mark read/delete.",
            "Blog public có slug unique và text index để hiển thị nội dung landing/knowledge sharing.",
        ],
    },
]


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Capstone Project Report")
    set_font(r, size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Report 4 - Software Design Document")
    set_font(r, size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Learning Management System (LMS)")
    set_font(r, size=18, bold=True, color=DARK_BLUE)
    doc.add_paragraph()
    meta = [
        ("Project source", "BE_LMS (ExpressJS/TypeScript) + FE_LMS (React/Vite)"),
        ("Document scope", "System design, package design, database design, detailed class/sequence design"),
        ("Prepared date", date.today().strftime("%Y-%m-%d")),
    ]
    add_matrix_table(doc, ["Item", "Value"], meta, [1.7, 4.8], font_size=10)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hanoi, July 2026")
    set_font(r, size=12, color=MUTED)
    doc.add_page_break()


def add_record_of_changes(doc: Document) -> None:
    doc.add_heading("I. Record of Changes", level=1)
    add_para(doc, "*A - Added    M - Modified    D - Deleted")
    rows = [
        ["2026-07-16", "A/M", "Codex", "Hoàn thiện toàn bộ nội dung SDD dựa trên source code LMS: system architecture, package diagrams, database design, detailed class specifications và sequence flows."],
    ]
    add_matrix_table(doc, ["Date", "A* M, D", "In charge", "Change Description"], rows, [1.15, 0.9, 1.15, 3.3], font_size=9)


def add_toc(doc: Document) -> None:
    doc.add_heading("II. Software Design Document", level=1)
    doc.add_heading("Document Navigation", level=2)
    rows = [
        ["1", "System Design", "Architecture, deployment view, backend/frontend package diagrams and explanations."],
        ["2", "Database Design", "ERD views, collection/table descriptions, primary keys, foreign keys and constraints."],
        ["3", "Detailed Design", "Class diagrams, class specifications and sequence diagrams for the main LMS flows."],
    ]
    add_matrix_table(doc, ["Section", "Name", "Coverage"], rows, [0.8, 1.8, 3.9], font_size=9)
    add_para(doc, "This document is authored from the actual project structure and code paths under BE_LMS and FE_LMS, especially app.ts, routes, controllers, services, models, socket handlers, validators, frontend AppRoutes, services and stores.")


def add_system_design(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("1. System Design", level=2)
    doc.add_heading("1.1 System Architecture", level=3)
    add_para(doc, "The LMS is implemented as a MERN-style learning platform with a React/Vite single-page frontend, an ExpressJS TypeScript backend, MongoDB persistence through Mongoose, MinIO/S3 object storage, Resend transactional email and Socket.IO realtime communication.")
    add_picture(doc, diagrams["architecture"], "Figure 1. System architecture for LMS runtime and external integrations.")
    rows = [
        ["React/Vite SPA", "FE_LMS", "Delivers pages, protected routes, dashboards and feature UI. Uses Axios services to call backend APIs and Socket.IO client for realtime chat/video."],
        ["Express API", "BE_LMS/src/app.ts", "Registers REST endpoints, CORS, cookies, custom response middleware, error handler, multer upload and Socket.IO server."],
        ["MongoDB/Mongoose", "BE_LMS/src/models", "Stores all business documents using 31 Mongoose models with refs and indexes."],
        ["MinIO/S3", "BE_LMS/src/config/minio.ts, utils/uploadFile.ts", "Stores binary files: lesson material, course logo, submissions, feedback attachments and chat files."],
        ["Resend", "BE_LMS/src/utils/sendMail.ts", "Sends email verification, password reset, invite and absence warning emails."],
        ["Socket.IO", "BE_LMS/src/socket", "Authenticates socket by JWT cookie, joins chat rooms, broadcasts messages and relays WebRTC signaling."],
        ["Google OAuth", "auth.service.ts", "Optional OAuth login; creates/updates verified users and sessions."],
    ]
    add_matrix_table(doc, ["Component", "Code Location", "Design Responsibility"], rows, [1.25, 1.8, 3.45], font_size=8)
    add_para(doc, "Key architectural rule: HTTP request processing follows the same chain across modules: route -> middleware -> controller -> validator -> service -> model/external integration -> custom response. Realtime events follow socket middleware -> event handler -> service/model -> room emission.")

    doc.add_heading("1.2 Package Diagram", level=3)
    add_para(doc, "The backend is organized as a layered architecture. routes and socket are the entry points, controllers coordinate request/response, services own business rules, models own persistence, while middleware/validators/utils/config/types/constants are cross-cutting support packages.")
    add_picture(doc, diagrams["backend_package"], "Figure 2. Backend package diagram.")
    add_matrix_table(doc, ["No", "Package", "Description"], BACKEND_PACKAGES, [0.45, 1.25, 4.8], font_size=8)
    add_para(doc, "The frontend separates navigation, feature pages, reusable components, service clients, state stores, hooks and shared type definitions. This keeps API integration and UI rendering independent while still matching the backend module names.")
    add_picture(doc, diagrams["frontend_package"], "Figure 3. Frontend package diagram.")
    add_matrix_table(doc, ["No", "Package", "Description"], FRONTEND_PACKAGES, [0.45, 1.25, 4.8], font_size=8)


def add_database_design(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("2. Database Design", level=2)
    add_para(doc, "The database layer uses MongoDB documents through Mongoose. All collections use MongoDB _id as the primary key. Relationships are represented through ObjectId refs, with indexes used for uniqueness, filtering and operational safety.")
    add_picture(doc, diagrams["erd_core"], "Figure 4. ERD for core academic domain.")
    add_picture(doc, diagrams["erd_learning"], "Figure 5. ERD for learning content and assessment.")
    add_picture(doc, diagrams["erd_schedule"], "Figure 6. ERD for schedule and attendance.")
    add_picture(doc, diagrams["erd_communication"], "Figure 7. ERD for communication and support modules.")

    doc.add_heading("Table Descriptions", level=3)
    add_matrix_table(doc, ["No", "Table / Collection", "Primary Key", "Foreign Keys", "Description"], TABLE_DESCRIPTIONS, [0.35, 1.15, 0.75, 1.75, 2.5], font_size=7)
    doc.add_heading("Important Constraints and Indexes", level=3)
    rows = [
        ["Enrollment", "Unique { studentId, courseId }", "Prevents duplicated enrollment records for the same student and course."],
        ["Attendance", "Unique { courseId, studentId, date }", "Allows exactly one attendance status per student per course session date."],
        ["Course", "Unique slug; indexes on isPublished/status/isDeleted", "Supports public listing, recycle bin and course detail lookup."],
        ["Subject", "Unique name/code/slug; prerequisite refs", "Protects catalog identity and prerequisite map."],
        ["ClassSchedule", "Unique active teacher-day-timeSlot partial index", "Prevents teacher double-booking when schedule is approved/active."],
        ["LessonProgress", "Unique { studentId, lessonId }", "Keeps one progress record per student per lesson."],
        ["QuizAttempt", "Unique { quizId, studentId }", "Controls a single attempt record per quiz/student in current design."],
        ["CourseInvite", "Unique tokenHash", "Protects invite tokens and allows revocation/expiry checks."],
    ]
    add_matrix_table(doc, ["Entity", "Constraint / Index", "Design Purpose"], rows, [1.3, 2.0, 3.2], font_size=8)


def add_feature(doc: Document, diagrams: dict[str, Path], feature: dict) -> None:
    key = feature["key"]
    doc.add_heading(feature["heading"], level=3)
    add_para(doc, feature["purpose"])
    add_picture(doc, diagrams[f"class_{key}"], f"Class Diagram - {feature['heading'].split(' ', 1)[1]}.")
    add_matrix_table(doc, ["Class / Module", "Specification"], feature["classes"], [1.75, 4.75], font_size=8)
    add_para(doc, "Main sequence and business flow:")
    for step in feature["flow"]:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(step)
        set_font(r)
    seq_key = {
        "auth": "seq_auth",
        "enrollment": "seq_enrollment",
        "lesson": "seq_lesson",
        "quiz": "seq_quiz",
        "assignment": "seq_assignment",
        "attendance": "seq_attendance",
        "chat": "seq_chat",
        "feedback": "seq_feedback",
    }[key]
    add_picture(doc, diagrams[seq_key], f"Sequence Diagram - {feature['heading'].split(' ', 1)[1]}.")


def add_detailed_design(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("3. Detailed Design", level=2)
    add_para(doc, "This section documents the principal LMS flows from UI trigger to persistence/external integration. Similar modules share the same layered pattern, so each feature presents the relevant classes, responsibilities and a representative sequence diagram.")
    for feature in FEATURES:
        add_feature(doc, diagrams, feature)

    doc.add_heading("3.9 Cross-Cutting Design Decisions", level=3)
    rows = [
        ["Authentication", "JWT accessToken and refreshToken are stored as HTTP-only cookies. refreshToken maps to SessionModel, not localStorage."],
        ["Authorization", "Protected routes use authenticate and authorize(Role...). Socket.IO uses socketAuthMiddleware to verify the same access token from cookies."],
        ["Validation", "Validators parse request DTOs before service logic; services still enforce domain rules and ownership checks."],
        ["Error handling", "Services use appAssert/AppError with HTTP constants; errorHandler centralizes failure responses."],
        ["File upload", "Multer memoryStorage enforces MIME allow-list and 20MB limit, then uploadFile streams buffer into MinIO/S3."],
        ["Soft delete", "Course and notification-like flows preserve history using isDeleted/deletedAt/deletedBy where applicable."],
        ["Realtime", "Socket events update MongoDB first, then emit to chat room/user rooms to keep UI state consistent."],
        ["Observability", "Service-level indexes and filtered queries support pagination, stats, text search and role-scoped access."],
    ]
    add_matrix_table(doc, ["Concern", "Design Decision"], rows, [1.45, 5.05], font_size=8)

    doc.add_heading("3.10 Traceability Summary", level=3)
    rows = [
        ["Authentication & Users", "auth.route.ts, auth.controller.ts, auth.service.ts, user/session/verificationCode models", "Register, verify email, login, Google login, refresh, logout, reset password, profile/user management."],
        ["Course & Enrollment", "course/enrollment/subject/courseInvite services and models", "Course creation, publish, soft delete, self-enroll, invite join, prerequisite, capacity and cooldown."],
        ["Learning Content", "lesson, lessonMaterial, lessonProgress modules", "Lesson CRUD, material upload/download, presigned URL, time-on-task and completion."],
        ["Assessment", "quiz, quizQuestion, quizAttempt, assignment, submission modules", "Question import/export, quiz attempt auto-save/submit/regrade/ban, assignment submission/grading/report."],
        ["Schedule & Attendance", "schedule, timeSlot, scheduleException, attendance modules", "Weekly schedule, exception approval, attendance marking, exports and absence email warning."],
        ["Communication", "forum, chatRoom, message, notification, announcement, feedback, blog, socket modules", "Forum threads/replies, realtime chat, WebRTC signaling, notifications, announcements, ratings and public posts."],
    ]
    add_matrix_table(doc, ["Business Area", "Code Evidence", "Covered Flows"], rows, [1.35, 2.15, 3.0], font_size=8)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = make_diagrams()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_record_of_changes(doc)
    add_toc(doc)
    add_system_design(doc, diagrams)
    add_database_design(doc, diagrams)
    add_detailed_design(doc, diagrams)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())
